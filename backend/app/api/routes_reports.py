import csv
import io
import json
import logging
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.api.deps import get_db, require_role
from app.db import models
from app.risk.cvss_engine import cvss_baseline_from_severity


def _cvss_or_baseline(f: models.Finding) -> float | None:
    """Use the stored CVSS, falling back to a severity-derived baseline so
    plugin findings without a CVE still get a representative score."""
    return f.cvss_base if f.cvss_base is not None else cvss_baseline_from_severity(f.severity)

logger = logging.getLogger("vulnscan.reports")

router = APIRouter(prefix="/reports", tags=["reports"])


# ─── Templates ──────────────────────────────────────────────────────────────
ALL_SECTIONS = [
    {"id": "summary",     "name": "Executive summary"},
    {"id": "severity",    "name": "Severity overview"},
    {"id": "findings",    "name": "Findings table"},
    {"id": "evidence",    "name": "Evidence (HTTP requests / payloads)"},
    {"id": "remediation", "name": "Remediation guidance"},
    {"id": "compliance",  "name": "Compliance mapping (NIST / PCI / CIS)"},
    {"id": "ai_summary",  "name": "AI-generated executive summary"},
]

BUILTIN_TEMPLATES = [
    {
        "name": "Executive + Technical",
        "description": "Default — exec summary, severity overview, full findings, remediation, AI analysis",
        "sections": ["summary", "severity", "findings", "remediation", "ai_summary"],
        "severities": ["critical", "high", "medium", "low", "info"],
        "options": {},
        "builtin": True,
    },
    {
        "name": "Executive Summary",
        "description": "Cover, severity overview, top findings only — no technical evidence",
        "sections": ["summary", "severity"],
        "severities": ["critical", "high"],
        "options": {"top_findings": 5},
        "builtin": True,
    },
    {
        "name": "Technical Detail",
        "description": "All findings with evidence, remediation steps, AI verdicts — no exec summary",
        "sections": ["severity", "findings", "evidence", "remediation", "ai_summary"],
        "severities": ["critical", "high", "medium", "low", "info"],
        "options": {},
        "builtin": True,
    },
    {
        "name": "PCI DSS Compliance",
        "description": "Findings mapped to PCI 4.0 requirements",
        "sections": ["summary", "severity", "findings", "compliance"],
        "severities": ["critical", "high", "medium"],
        "options": {"compliance_filter": "PCI"},
        "builtin": True,
    },
]


def _seed_builtins(db: Session, ws: int) -> None:
    """Ensure built-in templates exist for this workspace."""
    existing = {t.name for t in db.query(models.ReportTemplate).filter(
        models.ReportTemplate.workspace_id == ws,
        models.ReportTemplate.builtin == True,  # noqa: E712
    ).all()}
    for tpl in BUILTIN_TEMPLATES:
        if tpl["name"] in existing:
            continue
        db.add(models.ReportTemplate(
            workspace_id=ws,
            name=tpl["name"],
            description=tpl["description"],
            builtin=True,
            sections_json=json.dumps(tpl["sections"]),
            severities_json=json.dumps(tpl["severities"]),
            options_json=json.dumps(tpl["options"]),
        ))
    db.commit()


def _serialize_template(t: models.ReportTemplate) -> dict:
    return {
        "id": t.id,
        "name": t.name,
        "description": t.description,
        "builtin": t.builtin,
        "sections": json.loads(t.sections_json or "[]"),
        "severities": json.loads(t.severities_json or "[]"),
        "options": json.loads(t.options_json or "{}"),
        "created_at": t.created_at.isoformat() if t.created_at else None,
    }


class TemplateCreate(BaseModel):
    name: str
    description: str | None = ""
    sections: list[str] | None = None
    severities: list[str] | None = None
    options: dict | None = None


class TemplateUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    sections: list[str] | None = None
    severities: list[str] | None = None
    options: dict | None = None


@router.get("/templates/sections")
def list_sections(user=Depends(require_role("admin", "analyst", "viewer"))):
    return ALL_SECTIONS


@router.get("/templates")
def list_templates(
    user=Depends(require_role("admin", "analyst", "viewer")),
    db: Session = Depends(get_db),
):
    _seed_builtins(db, user["ws"])
    rows = (
        db.query(models.ReportTemplate)
        .filter(models.ReportTemplate.workspace_id == user["ws"])
        .order_by(models.ReportTemplate.builtin.desc(), models.ReportTemplate.name)
        .all()
    )
    return [_serialize_template(t) for t in rows]


@router.post("/templates")
def create_template(
    body: TemplateCreate,
    user=Depends(require_role("admin", "analyst")),
    db: Session = Depends(get_db),
):
    if not body.name or not body.name.strip():
        raise HTTPException(400, "name is required")
    valid = {s["id"] for s in ALL_SECTIONS}
    sections = [s for s in (body.sections or ["summary", "severity", "findings"]) if s in valid]
    if not sections:
        raise HTTPException(400, "At least one valid section is required")

    t = models.ReportTemplate(
        workspace_id=user["ws"],
        name=body.name.strip(),
        description=body.description or "",
        builtin=False,
        sections_json=json.dumps(sections),
        severities_json=json.dumps(body.severities or ["critical", "high", "medium", "low", "info"]),
        options_json=json.dumps(body.options or {}),
    )
    db.add(t)
    db.commit()
    db.refresh(t)
    return _serialize_template(t)


@router.patch("/templates/{tpl_id}")
def update_template(
    tpl_id: int,
    body: TemplateUpdate,
    user=Depends(require_role("admin", "analyst")),
    db: Session = Depends(get_db),
):
    t = db.query(models.ReportTemplate).filter(
        models.ReportTemplate.workspace_id == user["ws"],
        models.ReportTemplate.id == tpl_id,
    ).first()
    if not t:
        raise HTTPException(404, "Template not found")
    if t.builtin:
        raise HTTPException(400, "Built-in templates can't be edited — clone first")

    data = body.dict(exclude_unset=True)
    if "name" in data and data["name"]:
        t.name = data["name"].strip()
    if "description" in data:
        t.description = data["description"] or ""
    if "sections" in data:
        valid = {s["id"] for s in ALL_SECTIONS}
        clean = [s for s in (data["sections"] or []) if s in valid]
        if clean:
            t.sections_json = json.dumps(clean)
    if "severities" in data:
        t.severities_json = json.dumps(data["severities"] or [])
    if "options" in data:
        t.options_json = json.dumps(data["options"] or {})
    t.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(t)
    return _serialize_template(t)


@router.post("/templates/{tpl_id}/clone")
def clone_template(
    tpl_id: int,
    user=Depends(require_role("admin", "analyst")),
    db: Session = Depends(get_db),
):
    src = db.query(models.ReportTemplate).filter(
        models.ReportTemplate.workspace_id == user["ws"],
        models.ReportTemplate.id == tpl_id,
    ).first()
    if not src:
        raise HTTPException(404, "Template not found")
    t = models.ReportTemplate(
        workspace_id=user["ws"],
        name=f"{src.name} (copy)",
        description=src.description,
        builtin=False,
        sections_json=src.sections_json,
        severities_json=src.severities_json,
        options_json=src.options_json,
    )
    db.add(t)
    db.commit()
    db.refresh(t)
    return _serialize_template(t)


@router.delete("/templates/{tpl_id}")
def delete_template(
    tpl_id: int,
    user=Depends(require_role("admin", "analyst")),
    db: Session = Depends(get_db),
):
    t = db.query(models.ReportTemplate).filter(
        models.ReportTemplate.workspace_id == user["ws"],
        models.ReportTemplate.id == tpl_id,
    ).first()
    if not t:
        raise HTTPException(404, "Template not found")
    if t.builtin:
        raise HTTPException(400, "Built-in templates can't be deleted")
    db.delete(t)
    db.commit()
    return {"ok": True, "deleted_template_id": tpl_id}


# ─── Report list + generation ───────────────────────────────────────────────

def _job_findings(db: Session, ws: int, job_id: int):
    job = (
        db.query(models.ScanJob)
        .filter(models.ScanJob.workspace_id == ws, models.ScanJob.id == job_id)
        .first()
    )
    if not job:
        raise HTTPException(404, "scan job not found")
    findings = (
        db.query(models.Finding)
        .filter(models.Finding.workspace_id == ws, models.Finding.job_id == job_id)
        .order_by(models.Finding.risk_score.desc().nullslast())
        .all()
    )
    return job, findings


def _latest_ai_analysis(db: Session, ws: int, job_id: int) -> dict | None:
    """Return the latest done AI analysis for this job as a dict, or None."""
    a = (
        db.query(models.AiAnalysis)
        .filter(
            models.AiAnalysis.workspace_id == ws,
            models.AiAnalysis.job_id == job_id,
            models.AiAnalysis.status == "done",
        )
        .order_by(models.AiAnalysis.finished_at.desc().nullslast(),
                  models.AiAnalysis.id.desc())
        .first()
    )
    if not a or not a.result_json:
        return None
    try:
        result = json.loads(a.result_json)
    except Exception:
        return None
    return {
        "id": a.id,
        "provider": a.provider,
        "mode": a.mode,
        "finished_at": a.finished_at.isoformat() if a.finished_at else None,
        "duration_seconds": a.duration_seconds,
        "result": result,
    }


def _ai_lookup(ai: dict | None) -> dict[str, dict]:
    """Build {finding_id|fingerprint|title -> entry} so report generators can match per-finding."""
    if not ai or not ai.get("result"):
        return {}
    result = ai["result"]
    poc_map = result.get("poc_results") or {}
    out: dict[str, dict] = {}

    # Backend returns finding_validations as dict keyed by finding_id
    validations = result.get("finding_validations") or {}
    for fid, v in validations.items():
        entry = {"finding_id": fid, **(v if isinstance(v, dict) else {})}
        if fid in poc_map:
            entry["poc_code"] = poc_map[fid]
        out[f"id:{fid}"] = entry

    # Fallback: array format (findings / findings_validated)
    entries = result.get("findings") or result.get("findings_validated") or []
    for e in entries:
        if e.get("finding_id") is not None:
            fid = str(e["finding_id"])
            if fid not in out:
                if fid in poc_map:
                    e["poc_code"] = poc_map[fid]
                out[f"id:{fid}"] = e
        if e.get("fingerprint"):
            out[f"fp:{e['fingerprint']}"] = e
        if e.get("title"):
            out[f"t:{e['title']}"] = e
    return out


def _ai_match(lookup: dict, finding: models.Finding) -> dict | None:
    return (
        lookup.get(f"id:{finding.id}")
        or (finding.fingerprint and lookup.get(f"fp:{finding.fingerprint}"))
        or (finding.title and lookup.get(f"t:{finding.title}"))
        or None
    )


def _resolve_template(db: Session, ws: int, template_id: int | None) -> dict:
    """Load template config or fall back to a default that includes everything."""
    if template_id:
        t = db.query(models.ReportTemplate).filter(
            models.ReportTemplate.workspace_id == ws,
            models.ReportTemplate.id == template_id,
        ).first()
        if t:
            return _serialize_template(t)
    # Default — include everything
    return {
        "name": "Default",
        "sections": [s["id"] for s in ALL_SECTIONS if s["id"] != "ai_summary"],
        "severities": ["critical", "high", "medium", "low", "info"],
        "options": {},
    }


@router.get("")
def list_reports(
    user=Depends(require_role("admin", "analyst", "viewer")),
    db: Session = Depends(get_db),
):
    rows = (
        db.query(models.ScanJob)
        .filter(
            models.ScanJob.workspace_id == user["ws"],
            models.ScanJob.status == "done",
        )
        .order_by(models.ScanJob.id.desc())
        .all()
    )

    out = []
    for j in rows:
        sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
        f_rows = (
            db.query(models.Finding.severity)
            .filter(
                models.Finding.workspace_id == user["ws"],
                models.Finding.job_id == j.id,
            )
            .all()
        )
        total = 0
        for (sev,) in f_rows:
            total += 1
            if sev in sev_counts:
                sev_counts[sev] += 1

        asset_name = None
        if j.asset_id:
            a = db.query(models.Asset).filter(models.Asset.id == j.asset_id).first()
            asset_name = a.name if a else None

        out.append({
            "id": f"R-{j.id}",
            "job_id": j.id,
            "target": j.target,
            "scan_type": j.scan_type,
            "profile_id": j.profile_id,
            "asset_id": j.asset_id,
            "asset_name": asset_name,
            "findings": total,
            "severity_counts": sev_counts,
            "generated_at": j.finished_at.isoformat() if j.finished_at else (j.created_at.isoformat() if j.created_at else None),
            "formats": ["json", "csv", "sarif", "md", "pdf", "docx", "xlsx"],
        })

    return out


@router.get("/{job_id}/download")
def download_report(
    job_id: int,
    format: str = "json",
    template_id: int | None = None,
    user=Depends(require_role("admin", "analyst", "viewer")),
    db: Session = Depends(get_db),
):
    """Generate and stream the report for a scan job in the requested format."""
    fmt = format.lower()
    if fmt not in ("json", "csv", "sarif", "md", "pdf", "docx", "xlsx"):
        raise HTTPException(400, f"Unsupported format: {format}. Use json|csv|sarif|md|pdf|docx|xlsx")

    job, findings = _job_findings(db, user["ws"], job_id)
    template = _resolve_template(db, user["ws"], template_id)

    # Apply severity filter from template
    sev_filter = set(template.get("severities") or ["critical", "high", "medium", "low", "info"])
    findings = [f for f in findings if f.severity in sev_filter]

    # Fetch latest AI analysis if the template wants the ai_summary section
    ai = None
    if "ai_summary" in (template.get("sections") or []):
        ai = _latest_ai_analysis(db, user["ws"], job_id)

    fname_base = f"vulnscan-report-{job.id}-{datetime.now(timezone.utc).strftime('%Y%m%d')}"

    # Always fetch AI analysis for CSV/XLSX too (not just when template has ai_summary)
    if not ai:
        ai = _latest_ai_analysis(db, user["ws"], job_id)
    ai_lookup = _ai_lookup(ai)

    if fmt == "json":
        return _stream(_as_json(job, findings, template, ai), "application/json", f"{fname_base}.json")
    if fmt == "csv":
        return _stream(_as_csv(findings, ai_lookup), "text/csv", f"{fname_base}.csv")
    if fmt == "sarif":
        return _stream(_as_sarif(job, findings), "application/sarif+json", f"{fname_base}.sarif")
    if fmt == "md":
        return _stream(_as_markdown(job, findings, template, ai), "text/markdown", f"{fname_base}.md")
    try:
        if fmt == "pdf":
            return _stream_bytes(_as_pdf(job, findings, template, ai), "application/pdf", f"{fname_base}.pdf")
        if fmt == "docx":
            return _stream_bytes(
                _as_docx(job, findings, template, ai),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{fname_base}.docx",
            )
        if fmt == "xlsx":
            return _stream_bytes(
                _as_xlsx(job, findings, ai_lookup),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                f"{fname_base}.xlsx",
            )
    except ModuleNotFoundError as exc:
        logger.error("Missing report dependency: %s", exc)
        raise HTTPException(500, f"Report generator missing — install '{exc.name}' (rebuild backend image with new requirements.txt)")
    except Exception as exc:
        logger.exception("Report generation failed for job #%d format=%s", job_id, fmt)
        raise HTTPException(500, f"{type(exc).__name__}: {exc}")


def _stream(body: str, media_type: str, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(body.encode("utf-8")),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _stream_bytes(body: bytes, media_type: str, filename: str) -> StreamingResponse:
    return StreamingResponse(
        io.BytesIO(body),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Light-weight format generators ──────────────────────────────────────────

def _as_json(job: models.ScanJob, findings: list[models.Finding], template: dict, ai: dict | None = None) -> str:
    payload = {
        "report_id": f"R-{job.id}",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "template": template.get("name"),
        "job": {
            "id": job.id,
            "target": job.target,
            "scan_type": job.scan_type,
            "status": job.status,
            "profile_id": job.profile_id,
            "asset_id": job.asset_id,
            "started_at": job.created_at.isoformat() if job.created_at else None,
            "finished_at": job.finished_at.isoformat() if job.finished_at else None,
        },
        "findings": [
            {
                "id": f.id,
                "title": f.title,
                "severity": f.severity,
                "plugin_id": f.plugin_id,
                "target": f.target,
                "cvss_base": _cvss_or_baseline(f),
                "risk_score": f.risk_score,
                "confidence": f.confidence,
                "is_kev": f.is_kev,
                "sla_days": f.sla_days,
                "description": f.description,
                "remediation": f.remediation,
                "evidence": f.evidence,
                "fingerprint": f.fingerprint,
                "compliance": json.loads(f.compliance_json) if f.compliance_json else None,
                "references": json.loads(f.references_json) if f.references_json else [],
                "status": f.status,
            }
            for f in findings
        ],
    }
    if ai:
        payload["ai_analysis"] = ai
    return json.dumps(payload, indent=2)


def _as_csv(findings: list[models.Finding], ai_lookup: dict | None = None) -> str:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["id", "severity", "title", "target", "plugin_id", "cvss", "risk", "confidence",
                "is_kev", "sla_days", "fingerprint", "ai_verdict", "ai_confidence", "ai_reasoning"])
    for f in findings:
        cvss_val = _cvss_or_baseline(f)
        ai_entry = _ai_match(ai_lookup or {}, f)
        w.writerow([
            f.id, f.severity, f.title, f.target, f.plugin_id,
            cvss_val if cvss_val is not None else "",
            f.risk_score if f.risk_score is not None else "",
            f.confidence if f.confidence is not None else "",
            "yes" if f.is_kev else "no",
            f.sla_days if f.sla_days is not None else "",
            f.fingerprint or "",
            (ai_entry.get("verdict") or "") if ai_entry else "",
            (ai_entry.get("confidence") or "") if ai_entry else "",
            (ai_entry.get("reasoning") or "") if ai_entry else "",
        ])
    return buf.getvalue()


def _as_sarif(job: models.ScanJob, findings: list[models.Finding]) -> str:
    sev_to_level = {
        "critical": "error", "high": "error",
        "medium": "warning", "low": "note", "info": "note",
    }
    rules: dict[str, dict] = {}
    results = []
    for f in findings:
        rule_id = f.plugin_id or "unknown"
        if rule_id not in rules:
            rules[rule_id] = {
                "id": rule_id,
                "shortDescription": {"text": f.title[:120] if f.title else rule_id},
                "fullDescription":  {"text": f.description or f.title or rule_id},
                "defaultConfiguration": {"level": sev_to_level.get(f.severity, "warning")},
            }
        results.append({
            "ruleId": rule_id,
            "level": sev_to_level.get(f.severity, "warning"),
            "message": {"text": f.title or "Finding"},
            "properties": {
                "severity": f.severity,
                "cvss_base": _cvss_or_baseline(f),
                "risk_score": f.risk_score,
                "confidence": f.confidence,
                "is_kev": f.is_kev,
            },
            "locations": [{
                "physicalLocation": {"artifactLocation": {"uri": f.target or job.target}},
            }],
            "fingerprints": {"vulnscan/v1": f.fingerprint or ""},
        })

    sarif = {
        "version": "2.1.0",
        "$schema": "https://json.schemastore.org/sarif-2.1.0.json",
        "runs": [{
            "tool": {"driver": {
                "name": "VulnScan",
                "informationUri": "https://github.com/vulnscan/vulnscan",
                "rules": list(rules.values()),
            }},
            "results": results,
            "invocations": [{
                "executionSuccessful": job.status == "done",
                "endTimeUtc": job.finished_at.isoformat() if job.finished_at else None,
            }],
        }],
    }
    return json.dumps(sarif, indent=2)


def _as_markdown(job: models.ScanJob, findings: list[models.Finding], template: dict, ai: dict | None = None) -> str:
    sections = set(template.get("sections") or [])
    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        if f.severity in sev_counts:
            sev_counts[f.severity] += 1
    sev_emoji = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢", "info": "⚪"}
    ai_lookup = _ai_lookup(ai)

    lines = []
    lines.append(f"# VulnScan Report — Job #{job.id}")
    lines.append("")
    lines.append(f"**Target:** `{job.target}`  ")
    lines.append(f"**Template:** {template.get('name', 'Default')}  ")
    lines.append(f"**Scan type:** {job.scan_type}  ")
    if job.created_at:
        lines.append(f"**Started:** {job.created_at.isoformat()}  ")
    if job.finished_at:
        lines.append(f"**Finished:** {job.finished_at.isoformat()}  ")
    lines.append(f"**Generated:** {datetime.now(timezone.utc).isoformat()}")
    lines.append("")

    if "summary" in sections:
        lines.append("## Executive Summary")
        lines.append("")
        n_total = len(findings)
        n_critical = sev_counts["critical"]
        n_high = sev_counts["high"]
        risk_label = "elevated" if n_critical > 0 else "moderate" if n_high > 0 else "manageable"
        lines.append(
            f"This scan of **{job.target}** identified **{n_total}** findings "
            f"across {len([s for s in sev_counts.values() if s > 0])} severity levels. "
            f"The current risk posture is **{risk_label}** — "
            f"{n_critical} critical, {n_high} high, {sev_counts['medium']} medium, "
            f"{sev_counts['low']} low, {sev_counts['info']} informational."
        )
        lines.append("")

    if "ai_summary" in sections and ai:
        lines.append("## AI Analysis")
        lines.append("")
        ai_result = ai.get("result") or {}
        ai_summary = ai_result.get("summary") or ai_result.get("executive_summary")
        if ai_summary:
            lines.append(str(ai_summary))
            lines.append("")
        chains = ai_result.get("attack_chains") or ai_result.get("chains") or []
        if chains:
            lines.append("### Attack chains")
            lines.append("")
            for i, c in enumerate(chains, 1):
                lines.append(f"**{i}. {c.get('name', c.get('title', f'Chain {i}'))}**")
                if c.get("description"):
                    lines.append("")
                    lines.append(c["description"])
                steps = c.get("steps") or c.get("path") or []
                if steps:
                    lines.append("")
                    for j, s in enumerate(steps, 1):
                        lines.append(f"  {j}. {s if isinstance(s, str) else s.get('text', json.dumps(s))}")
                lines.append("")
        recs = ai_result.get("recommendations") or ai_result.get("next_steps") or ai_result.get("remediation_priority") or []
        if recs:
            lines.append("### Remediation Priority")
            lines.append("")
            for r in recs:
                if isinstance(r, str):
                    lines.append(f"- {r}")
                else:
                    rank = f"**#{r['rank']}** " if r.get("rank") else ""
                    action = r.get("action") or r.get("text") or json.dumps(r)
                    meta = []
                    if r.get("timeframe"): meta.append(r["timeframe"])
                    if r.get("effort"): meta.append(f"effort: {r['effort']}")
                    if r.get("impact"): meta.append(f"impact: {r['impact']}")
                    suffix = f" ({', '.join(meta)})" if meta else ""
                    lines.append(f"- {rank}{action}{suffix}")
            lines.append("")
        lines.append(f"_From AI analysis #{ai.get('id')} · {ai.get('provider')} · {ai.get('mode')}_")
        lines.append("")

    if "severity" in sections:
        lines.append("## Severity Overview")
        lines.append("")
        for sev, n in sev_counts.items():
            if n > 0:
                lines.append(f"- {sev_emoji[sev]} **{sev.capitalize()}:** {n}")
        lines.append("")

    if "findings" in sections:
        lines.append("## Findings")
        lines.append("")
        for i, f in enumerate(findings, 1):
            lines.append(f"### {i}. {sev_emoji.get(f.severity, '')} {f.title or f.plugin_id}")
            lines.append("")
            lines.append(f"- **Severity:** {f.severity}")
            if f.risk_score is not None:
                lines.append(f"- **Risk score:** {f.risk_score:.1f}")
            cvss_val = _cvss_or_baseline(f)
            if cvss_val is not None:
                cvss_label = "CVSS base" if f.cvss_base is not None else "CVSS base (severity-derived)"
                lines.append(f"- **{cvss_label}:** {cvss_val}")
            if f.is_kev:
                lines.append(f"- **CISA KEV:** yes")
            if f.target:
                lines.append(f"- **Target:** `{f.target}`")
            lines.append(f"- **Plugin:** `{f.plugin_id}`")
            if f.description:
                lines.append("")
                lines.append("**Description**")
                lines.append("")
                lines.append(f.description)
            if "evidence" in sections and f.evidence:
                lines.append("")
                lines.append("**Evidence**")
                lines.append("")
                lines.append("```")
                lines.append(f.evidence)
                lines.append("```")
            if "remediation" in sections and f.remediation:
                lines.append("")
                lines.append("**Remediation**")
                lines.append("")
                lines.append(f.remediation)
            if "compliance" in sections and f.compliance_json:
                try:
                    refs = json.loads(f.compliance_json)
                    if refs:
                        lines.append("")
                        lines.append(f"**Compliance:** {', '.join(refs)}")
                except Exception:
                    pass

            ai_entry = _ai_match(ai_lookup, f) if ai_lookup else None
            if ai_entry:
                lines.append("")
                lines.append("**AI verdict**")
                lines.append("")
                verdict = ai_entry.get("verdict") or ai_entry.get("assessment") or "—"
                conf = ai_entry.get("confidence")
                conf_str = f" ({int(conf*100)}% confidence)" if conf is not None else ""
                lines.append(f"- {str(verdict).replace('_', ' ')}{conf_str}")
                if ai_entry.get("reasoning") or ai_entry.get("explanation"):
                    lines.append("")
                    lines.append(ai_entry.get("reasoning") or ai_entry.get("explanation"))
                if ai_entry.get("recommendation"):
                    lines.append("")
                    lines.append(f"**AI recommendation:** {ai_entry['recommendation']}")
                # PoC from poc_results dict or inline fields
                poc_obj = ai_entry.get("poc_code")
                poc_text = None
                if poc_obj and isinstance(poc_obj, dict):
                    poc_text = poc_obj.get("code") or json.dumps(poc_obj, indent=2)
                else:
                    raw = ai_entry.get("exploit_code") or ai_entry.get("poc") or ai_entry.get("payload")
                    if raw:
                        poc_text = raw if isinstance(raw, str) else json.dumps(raw, indent=2)
                if poc_text:
                    lines.append("")
                    lines.append("**AI-generated PoC**")
                    if poc_obj and poc_obj.get("description"):
                        lines.append("")
                        lines.append(f"_{poc_obj['description']}_")
                    lines.append("")
                    lang = poc_obj.get("language", "") if poc_obj else ""
                    lines.append(f"```{lang}")
                    lines.append(poc_text)
                    lines.append("```")

            lines.append("")
            lines.append("---")
            lines.append("")

    return "\n".join(lines)


# ─── Heavyweight format generators ───────────────────────────────────────────

def _sev_color(sev: str) -> tuple[float, float, float]:
    """RGB 0–1 for ReportLab."""
    return {
        "critical": (0.94, 0.33, 0.43),
        "high":     (0.91, 0.50, 0.29),
        "medium":   (0.83, 0.66, 0.22),
        "low":      (0.37, 0.72, 0.48),
        "info":     (0.42, 0.45, 0.51),
    }.get(sev, (0.4, 0.45, 0.51))


def _esc(s) -> str:
    """XML-escape a value for ReportLab's Paragraph mini-language.
    Without this, any <, >, or & in finding titles / evidence / targets crashes the PDF parser.
    """
    from xml.sax.saxutils import escape as _xml_escape
    if s is None:
        return ""
    return _xml_escape(str(s), {"\"": "&quot;", "'": "&apos;"})


def _as_pdf(job: models.ScanJob, findings: list[models.Finding], template: dict, ai: dict | None = None) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether,
    )

    sections = set(template.get("sections") or [])
    ai_lookup = _ai_lookup(ai)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=f"VulnScan Report #{job.id}",
    )

    styles = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=20, leading=24, textColor=colors.HexColor("#0d1117"), spaceAfter=8)
    H2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, leading=18, textColor=colors.HexColor("#0d1117"), spaceBefore=14, spaceAfter=6)
    H3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11.5, leading=15, textColor=colors.HexColor("#0d1117"), spaceBefore=10, spaceAfter=4)
    Body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9.5, leading=13.5, textColor=colors.HexColor("#2a313c"), alignment=TA_LEFT)
    Code = ParagraphStyle("Code", parent=styles["BodyText"], fontName="Courier", fontSize=8.5, leading=11, textColor=colors.HexColor("#2a313c"), backColor=colors.HexColor("#f4f6f9"), borderPadding=4)
    Meta = ParagraphStyle("Meta", parent=styles["BodyText"], fontSize=8.5, leading=11, textColor=colors.HexColor("#6e7785"))

    flow = []

    # Cover
    flow.append(Paragraph("VulnScan — Security Scan Report", H1))
    flow.append(Paragraph(f"Job <b>#{job.id}</b> · Target <font face='Courier'>{_esc(job.target)}</font>", Body))
    flow.append(Spacer(1, 6))
    meta_lines = [
        f"<b>Template:</b> {_esc(template.get('name', 'Default'))}",
        f"<b>Scan type:</b> {_esc(job.scan_type)}",
        f"<b>Status:</b> {_esc(job.status)}",
        f"<b>Started:</b> {_esc(job.created_at.isoformat() if job.created_at else '—')}",
        f"<b>Finished:</b> {_esc(job.finished_at.isoformat() if job.finished_at else '—')}",
        f"<b>Generated:</b> {_esc(datetime.now(timezone.utc).isoformat())}",
    ]
    for ln in meta_lines:
        flow.append(Paragraph(ln, Meta))

    flow.append(Spacer(1, 12))

    # Severity counts
    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        if f.severity in sev_counts:
            sev_counts[f.severity] += 1

    if "summary" in sections:
        flow.append(Paragraph("Executive Summary", H2))
        risk_label = "elevated" if sev_counts["critical"] > 0 else "moderate" if sev_counts["high"] > 0 else "manageable"
        flow.append(Paragraph(
            f"This scan of <font face='Courier'>{_esc(job.target)}</font> identified "
            f"<b>{len(findings)}</b> findings across {len([v for v in sev_counts.values() if v > 0])} severity levels. "
            f"The current risk posture is <b>{risk_label}</b> — "
            f"{sev_counts['critical']} critical, {sev_counts['high']} high, {sev_counts['medium']} medium, "
            f"{sev_counts['low']} low, {sev_counts['info']} informational.",
            Body,
        ))

    if "severity" in sections:
        flow.append(Paragraph("Severity Overview", H2))
        sev_data = [["Severity", "Count"]]
        for sev in ("critical", "high", "medium", "low", "info"):
            if sev_counts[sev] > 0:
                sev_data.append([sev.capitalize(), str(sev_counts[sev])])
        if len(sev_data) > 1:
            sev_table = Table(sev_data, colWidths=[60 * mm, 30 * mm], hAlign="LEFT")
            ts = [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16191d")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING",    (0, 0), (-1, -1), 6),
                ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#d2d7df")),
            ]
            for i, sev in enumerate(["critical", "high", "medium", "low", "info"]):
                if sev_counts[sev] > 0:
                    r, g, b = _sev_color(sev)
                    row_idx = next((idx + 1 for idx, row in enumerate(sev_data[1:]) if row[0].lower() == sev), None)
                    if row_idx:
                        ts.append(("TEXTCOLOR", (0, row_idx), (0, row_idx), colors.Color(r, g, b)))
            sev_table.setStyle(TableStyle(ts))
            flow.append(sev_table)

    # AI Analysis section
    if "ai_summary" in sections and ai:
        flow.append(Paragraph("AI Analysis", H2))
        ai_result = ai.get("result") or {}

        # Summary
        ai_summary = ai_result.get("summary") or ai_result.get("executive_summary")
        if ai_summary:
            flow.append(Paragraph(_esc(str(ai_summary)).replace("\n", "<br/>"), Body))
            flow.append(Spacer(1, 8))

        # Attack chains
        chains = ai_result.get("attack_chains") or ai_result.get("chains") or []
        if chains:
            flow.append(Paragraph("<b>Attack chains identified</b>", Body))
            for i, c in enumerate(chains, 1):
                name = c.get("name") or c.get("title") or f"Chain {i}"
                flow.append(Paragraph(f"<b>{i}. {_esc(name)}</b>", Body))
                if c.get("description"):
                    flow.append(Paragraph(_esc(c["description"]), Meta))
                steps = c.get("steps") or c.get("path") or []
                for j, s in enumerate(steps, 1):
                    s_text = s if isinstance(s, str) else (s.get("text") or json.dumps(s))
                    flow.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{j}. {_esc(s_text)}", Meta))
                flow.append(Spacer(1, 4))

        # Recommendations / Remediation priority
        recs = ai_result.get("recommendations") or ai_result.get("next_steps") or ai_result.get("remediation_priority") or []
        if recs:
            flow.append(Paragraph("<b>Remediation Priority</b>", Body))
            for r in recs:
                if isinstance(r, str):
                    flow.append(Paragraph(f"• {_esc(r)}", Meta))
                else:
                    rank = f"#{r['rank']} " if r.get("rank") else ""
                    action = r.get("action") or r.get("text") or json.dumps(r)
                    meta_parts = []
                    if r.get("timeframe"): meta_parts.append(r["timeframe"])
                    if r.get("effort"): meta_parts.append(f"effort: {r['effort']}")
                    if r.get("impact"): meta_parts.append(f"impact: {r['impact']}")
                    suffix = f" <i>({', '.join(meta_parts)})</i>" if meta_parts else ""
                    flow.append(Paragraph(f"• <b>{_esc(rank)}</b>{_esc(action)}{suffix}", Meta))

        flow.append(Spacer(1, 6))
        flow.append(Paragraph(
            f"<i>From AI analysis #{ai.get('id')} · {_esc(ai.get('provider', ''))} · {_esc(ai.get('mode', ''))}"
            + (f" · {ai.get('duration_seconds', 0):.1f}s" if ai.get('duration_seconds') else "")
            + "</i>",
            Meta,
        ))

    # Findings
    if "findings" in sections:
        flow.append(Paragraph("Findings", H2))
        if not findings:
            flow.append(Paragraph("No findings recorded.", Body))
        for i, f in enumerate(findings, 1):
            block = []
            r, g, b = _sev_color(f.severity)
            sev_chip = f"<font color='#ffffff' backColor='{colors.Color(r, g, b).hexval()}'>&nbsp;{_esc(f.severity).upper()}&nbsp;</font>"
            kev_chip = " <b><font color='#f0556d'>[KEV]</font></b>" if f.is_kev else ""
            title = _esc(f.title or f.plugin_id)
            block.append(Paragraph(f"{i}. {sev_chip}{kev_chip} {title}", H3))

            attrs = []
            if f.risk_score is not None: attrs.append(f"<b>Risk:</b> {f.risk_score:.1f}")
            cvss_val = _cvss_or_baseline(f)
            if cvss_val is not None:
                # Mark severity-derived scores so readers don't confuse them with NVD-vetted ones
                cvss_label = "CVSS" if f.cvss_base is not None else "CVSS (sev-derived)"
                attrs.append(f"<b>{cvss_label}:</b> {cvss_val}")
            if f.confidence is not None: attrs.append(f"<b>Confidence:</b> {int(f.confidence * 100)}%")
            attrs.append(f"<b>Plugin:</b> <font face='Courier'>{_esc(f.plugin_id)}</font>")
            if f.target:
                attrs.append(f"<b>Target:</b> <font face='Courier'>{_esc(f.target)}</font>")
            block.append(Paragraph(" &nbsp;·&nbsp; ".join(attrs), Meta))

            if f.description:
                block.append(Spacer(1, 4))
                block.append(Paragraph(_esc(f.description).replace("\n", "<br/>"), Body))

            if "evidence" in sections and f.evidence:
                block.append(Spacer(1, 4))
                block.append(Paragraph("<b>Evidence</b>", Body))
                # Truncate very long evidence to keep PDF readable, escape, then format
                raw = f.evidence[:1500] + ("…" if len(f.evidence) > 1500 else "")
                ev = _esc(raw).replace("\n", "<br/>").replace(" ", "&nbsp;")
                block.append(Paragraph(ev, Code))

            if "remediation" in sections and f.remediation:
                block.append(Spacer(1, 4))
                block.append(Paragraph("<b>Remediation</b>", Body))
                block.append(Paragraph(_esc(f.remediation).replace("\n", "<br/>"), Body))

            if "compliance" in sections and f.compliance_json:
                try:
                    refs = json.loads(f.compliance_json)
                    if refs:
                        block.append(Spacer(1, 4))
                        block.append(Paragraph(f"<b>Compliance:</b> {_esc(', '.join(refs))}", Meta))
                except Exception:
                    pass

            # Per-finding AI verdict (only if ai_summary section is in template)
            if "ai_summary" in sections and ai_lookup:
                ai_entry = _ai_match(ai_lookup, f)
                if ai_entry:
                    block.append(Spacer(1, 6))
                    verdict = str(ai_entry.get("verdict") or ai_entry.get("assessment") or "—").replace("_", " ")
                    conf = ai_entry.get("confidence")
                    conf_str = f" ({int(conf*100)}%)" if conf is not None else ""
                    block.append(Paragraph(f"<b>AI verdict:</b> {_esc(verdict)}{conf_str}", Meta))
                    reasoning = ai_entry.get("reasoning") or ai_entry.get("explanation")
                    if reasoning:
                        block.append(Paragraph(_esc(reasoning).replace("\n", "<br/>"), Meta))
                    if ai_entry.get("recommendation"):
                        block.append(Paragraph(f"<b>AI recommendation:</b> {_esc(ai_entry['recommendation'])}", Meta))
                    poc_obj = ai_entry.get("poc_code")
                    poc_text = None
                    if poc_obj and isinstance(poc_obj, dict):
                        poc_text = poc_obj.get("code") or json.dumps(poc_obj, indent=2)
                    else:
                        raw = ai_entry.get("exploit_code") or ai_entry.get("poc") or ai_entry.get("payload")
                        if raw:
                            poc_text = raw if isinstance(raw, str) else json.dumps(raw, indent=2)
                    if poc_text:
                        label = "AI-generated PoC"
                        if poc_obj and poc_obj.get("language"):
                            label += f" ({_esc(poc_obj['language'])})"
                        block.append(Paragraph(f"<b>{label}</b>", Meta))
                        if poc_obj and poc_obj.get("description"):
                            block.append(Paragraph(f"<i>{_esc(poc_obj['description'])}</i>", Meta))
                        poc_text = poc_text[:1500] + ("…" if len(poc_text) > 1500 else "")
                        block.append(Paragraph(_esc(poc_text).replace("\n", "<br/>").replace(" ", "&nbsp;"), Code))

            block.append(Spacer(1, 10))

            try:
                flow.append(KeepTogether(block))
            except Exception as exc:
                # Defensive: if a finding still trips ReportLab, log and skip rather than 500 the whole report
                logger.warning("PDF: skipped finding #%d due to %s: %.100s", f.id, type(exc).__name__, str(exc))
                flow.append(Paragraph(f"<i>(Finding #{f.id} could not be rendered — see JSON export for details)</i>", Meta))

    doc.build(flow)
    return buf.getvalue()


def _as_docx(job: models.ScanJob, findings: list[models.Finding], template: dict, ai: dict | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sections = set(template.get("sections") or [])
    ai_lookup = _ai_lookup(ai)
    doc = Document()

    title = doc.add_heading(f"VulnScan Security Scan Report", level=0)

    p = doc.add_paragraph()
    p.add_run(f"Job #{job.id} · Target: ").italic = True
    p.add_run(job.target).bold = True

    meta = doc.add_paragraph()
    for label, val in [
        ("Template", template.get("name", "Default")),
        ("Scan type", job.scan_type),
        ("Status", job.status),
        ("Started", job.created_at.isoformat() if job.created_at else "—"),
        ("Finished", job.finished_at.isoformat() if job.finished_at else "—"),
        ("Generated", datetime.now(timezone.utc).isoformat()),
    ]:
        run = meta.add_run(f"{label}: ")
        run.bold = True
        meta.add_run(f"{val}\n")

    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        if f.severity in sev_counts:
            sev_counts[f.severity] += 1

    if "summary" in sections:
        doc.add_heading("Executive Summary", level=1)
        risk_label = "elevated" if sev_counts["critical"] > 0 else "moderate" if sev_counts["high"] > 0 else "manageable"
        doc.add_paragraph(
            f"This scan of {job.target} identified {len(findings)} findings. "
            f"Current risk posture is {risk_label} — "
            f"{sev_counts['critical']} critical, {sev_counts['high']} high, "
            f"{sev_counts['medium']} medium, {sev_counts['low']} low, "
            f"{sev_counts['info']} informational."
        )

    if "severity" in sections:
        doc.add_heading("Severity Overview", level=1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Severity"
        hdr[1].text = "Count"
        for sev in ("critical", "high", "medium", "low", "info"):
            if sev_counts[sev] > 0:
                row = tbl.add_row().cells
                row[0].text = sev.capitalize()
                row[1].text = str(sev_counts[sev])

    if "ai_summary" in sections and ai:
        doc.add_heading("AI Analysis", level=1)
        ai_result = ai.get("result") or {}
        ai_summary = ai_result.get("summary") or ai_result.get("executive_summary")
        if ai_summary:
            doc.add_paragraph(str(ai_summary))

        chains = ai_result.get("attack_chains") or ai_result.get("chains") or []
        if chains:
            doc.add_heading("Attack chains identified", level=2)
            for i, c in enumerate(chains, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {c.get('name') or c.get('title') or f'Chain {i}'}")
                run.bold = True
                if c.get("description"):
                    doc.add_paragraph(c["description"])
                steps = c.get("steps") or c.get("path") or []
                for s in steps:
                    s_text = s if isinstance(s, str) else (s.get("text") or json.dumps(s))
                    doc.add_paragraph(s_text, style="List Number")

        recs = ai_result.get("recommendations") or ai_result.get("next_steps") or ai_result.get("remediation_priority") or []
        if recs:
            doc.add_heading("Remediation Priority", level=2)
            for r in recs:
                if isinstance(r, str):
                    doc.add_paragraph(r, style="List Bullet")
                else:
                    rank = f"#{r['rank']} " if r.get("rank") else ""
                    action = r.get("action") or r.get("text") or json.dumps(r)
                    meta_parts = []
                    if r.get("timeframe"): meta_parts.append(r["timeframe"])
                    if r.get("effort"): meta_parts.append(f"effort: {r['effort']}")
                    if r.get("impact"): meta_parts.append(f"impact: {r['impact']}")
                    suffix = f" ({', '.join(meta_parts)})" if meta_parts else ""
                    doc.add_paragraph(f"{rank}{action}{suffix}", style="List Bullet")

        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"From AI analysis #{ai.get('id')} · {ai.get('provider')} · {ai.get('mode')}"
            + (f" · {ai.get('duration_seconds', 0):.1f}s" if ai.get('duration_seconds') else "")
        )
        meta_run.italic = True
        meta_run.font.size = Pt(9)

    if "findings" in sections:
        doc.add_heading("Findings", level=1)
        if not findings:
            doc.add_paragraph("No findings recorded.")
        for i, f in enumerate(findings, 1):
            h = doc.add_heading(f"{i}. [{f.severity.upper()}] {f.title or f.plugin_id}", level=2)
            r, g, b = _sev_color(f.severity)
            for run in h.runs:
                run.font.color.rgb = RGBColor(int(r * 255), int(g * 255), int(b * 255))
            if f.is_kev:
                p = doc.add_paragraph()
                kev = p.add_run("⚠ CISA KEV — actively exploited in the wild")
                kev.bold = True
                kev.font.color.rgb = RGBColor(240, 85, 109)

            attrs = doc.add_paragraph()
            cvss_val = _cvss_or_baseline(f)
            cvss_label = "CVSS base" if f.cvss_base is not None else "CVSS base (sev-derived)"
            for label, val in [
                ("Plugin", f.plugin_id),
                ("Target", f.target or job.target),
                ("Risk score", f"{f.risk_score:.1f}" if f.risk_score is not None else "—"),
                (cvss_label, str(cvss_val) if cvss_val is not None else "—"),
                ("Confidence", f"{int(f.confidence * 100)}%" if f.confidence is not None else "—"),
            ]:
                run = attrs.add_run(f"{label}: ")
                run.bold = True
                attrs.add_run(f"{val}    ")

            if f.description:
                doc.add_paragraph(f.description)
            if "evidence" in sections and f.evidence:
                doc.add_paragraph("Evidence:").runs[0].bold = True
                ev = doc.add_paragraph(f.evidence[:2000] + ("…" if len(f.evidence) > 2000 else ""))
                ev.runs[0].font.name = "Courier New"
                ev.runs[0].font.size = Pt(9)
            if "remediation" in sections and f.remediation:
                doc.add_paragraph("Remediation:").runs[0].bold = True
                doc.add_paragraph(f.remediation)
            if "compliance" in sections and f.compliance_json:
                try:
                    refs = json.loads(f.compliance_json)
                    if refs:
                        p = doc.add_paragraph()
                        p.add_run("Compliance: ").bold = True
                        p.add_run(", ".join(refs))
                except Exception:
                    pass

            if "ai_summary" in sections and ai_lookup:
                ai_entry = _ai_match(ai_lookup, f)
                if ai_entry:
                    p = doc.add_paragraph()
                    p.add_run("AI verdict: ").bold = True
                    verdict = str(ai_entry.get("verdict") or ai_entry.get("assessment") or "—").replace("_", " ")
                    conf = ai_entry.get("confidence")
                    p.add_run(verdict + (f" ({int(conf*100)}% confidence)" if conf is not None else ""))

                    reasoning = ai_entry.get("reasoning") or ai_entry.get("explanation")
                    if reasoning:
                        doc.add_paragraph(reasoning)

                    if ai_entry.get("recommendation"):
                        p = doc.add_paragraph()
                        p.add_run("AI recommendation: ").bold = True
                        p.add_run(ai_entry["recommendation"])

                    poc_obj = ai_entry.get("poc_code")
                    poc_text = None
                    if poc_obj and isinstance(poc_obj, dict):
                        poc_text = poc_obj.get("code") or json.dumps(poc_obj, indent=2)
                    else:
                        raw = ai_entry.get("exploit_code") or ai_entry.get("poc") or ai_entry.get("payload")
                        if raw:
                            poc_text = raw if isinstance(raw, str) else json.dumps(raw, indent=2)
                    if poc_text:
                        p = doc.add_paragraph()
                        label = "AI-generated PoC"
                        if poc_obj and poc_obj.get("language"):
                            label += f" ({poc_obj['language']})"
                        p.add_run(f"{label}:").bold = True
                        if poc_obj and poc_obj.get("description"):
                            desc_p = doc.add_paragraph()
                            desc_run = desc_p.add_run(poc_obj["description"])
                            desc_run.italic = True
                            desc_run.font.size = Pt(9)
                        ev = doc.add_paragraph(poc_text[:2500] + ("…" if len(poc_text) > 2500 else ""))
                        if ev.runs:
                            ev.runs[0].font.name = "Courier New"
                            ev.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _as_xlsx(job: models.ScanJob, findings: list[models.Finding], ai_lookup: dict | None = None) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

    wb = Workbook()

    # Summary sheet
    ws = wb.active
    ws.title = "Summary"
    ws["A1"] = f"VulnScan Report — Job #{job.id}"
    ws["A1"].font = Font(bold=True, size=16)
    ws.merge_cells("A1:D1")

    meta_rows = [
        ("Target", job.target),
        ("Scan type", job.scan_type),
        ("Status", job.status),
        ("Started", job.created_at.isoformat() if job.created_at else ""),
        ("Finished", job.finished_at.isoformat() if job.finished_at else ""),
        ("Generated", datetime.now(timezone.utc).isoformat()),
        ("Total findings", len(findings)),
    ]
    for i, (k, v) in enumerate(meta_rows, start=3):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws.cell(row=i, column=2, value=v)

    # Severity counts
    sev_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        if f.severity in sev_counts:
            sev_counts[f.severity] += 1

    sev_start = 3 + len(meta_rows) + 2
    ws.cell(row=sev_start - 1, column=1, value="Severity breakdown").font = Font(bold=True, size=12)
    for i, (sev, n) in enumerate(sev_counts.items(), start=sev_start):
        c1 = ws.cell(row=i, column=1, value=sev.capitalize())
        c2 = ws.cell(row=i, column=2, value=n)
        c1.font = Font(bold=True, color={
            "critical": "F0556D", "high": "E8804A", "medium": "D4A838",
            "low": "5FB87A", "info": "6B7382"
        }.get(sev, "000000"))

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 50

    # All-findings sheet + per-severity sheets
    def _write_finding_sheet(sheet, rows: list[models.Finding]):
        headers = ["ID", "Severity", "Title", "Target", "Plugin", "Risk", "CVSS", "Confidence",
                   "KEV", "SLA days", "Fingerprint", "AI Verdict", "AI Confidence", "AI Reasoning"]
        for col, h in enumerate(headers, start=1):
            c = sheet.cell(row=1, column=col, value=h)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill("solid", fgColor="16191D")
            c.alignment = Alignment(vertical="center")
        for r, f in enumerate(rows, start=2):
            sheet.cell(row=r, column=1, value=f.id)
            sev_cell = sheet.cell(row=r, column=2, value=f.severity)
            sev_cell.font = Font(bold=True, color={
                "critical": "F0556D", "high": "E8804A", "medium": "D4A838",
                "low": "5FB87A", "info": "6B7382"
            }.get(f.severity, "000000"))
            sheet.cell(row=r, column=3, value=f.title or "")
            sheet.cell(row=r, column=4, value=f.target or "")
            sheet.cell(row=r, column=5, value=f.plugin_id)
            sheet.cell(row=r, column=6, value=f.risk_score)
            sheet.cell(row=r, column=7, value=_cvss_or_baseline(f))
            sheet.cell(row=r, column=8, value=f.confidence)
            sheet.cell(row=r, column=9, value="yes" if f.is_kev else "no")
            sheet.cell(row=r, column=10, value=f.sla_days)
            sheet.cell(row=r, column=11, value=f.fingerprint or "")
            ai_entry = _ai_match(ai_lookup or {}, f)
            sheet.cell(row=r, column=12, value=(ai_entry.get("verdict") or "") if ai_entry else "")
            sheet.cell(row=r, column=13, value=ai_entry.get("confidence") if ai_entry else None)
            sheet.cell(row=r, column=14, value=(ai_entry.get("reasoning") or "") if ai_entry else "")
        widths = [8, 12, 60, 32, 28, 8, 8, 12, 6, 10, 32, 16, 14, 50]
        for i, w in enumerate(widths, start=1):
            col_letter = chr(64 + i) if i <= 26 else chr(64 + (i - 1) // 26) + chr(65 + (i - 1) % 26)
            sheet.column_dimensions[col_letter].width = w
        sheet.freeze_panes = "A2"

    all_sheet = wb.create_sheet("All findings")
    _write_finding_sheet(all_sheet, findings)

    # Per-severity sheets, only if non-empty
    for sev in ("critical", "high", "medium", "low", "info"):
        rows = [f for f in findings if f.severity == sev]
        if not rows:
            continue
        s = wb.create_sheet(sev.capitalize())
        _write_finding_sheet(s, rows)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
